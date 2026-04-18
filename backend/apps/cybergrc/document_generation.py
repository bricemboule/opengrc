import json
from datetime import date, datetime
from io import BytesIO


EXCLUDED_ROW_KEYS = {
    "id",
    "organization",
    "organization_id",
    "created_at",
    "updated_at",
    "deleted_at",
    "created_by",
    "updated_by",
    "created_by_email",
    "updated_by_email",
}

TEXT_CONTENT_TYPES = {
    "markdown": "text/markdown; charset=utf-8",
    "text": "text/plain; charset=utf-8",
    "json": "application/json; charset=utf-8",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def humanize_token(value):
    return str(value or "").replace("_", " ").replace("-", " ").strip().title()


def format_scalar(value):
    if value is None or value == "":
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="minutes")
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, list):
        return ", ".join(part for part in [format_scalar(item) for item in value] if part)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def choose_row_title(row, fallback_index):
    for key in ("title", "name", "code", "record_title"):
        label = format_scalar(row.get(key))
        if label:
            return label
    return f"Entry {fallback_index}"


def build_document_summary(module_label, rows_count, document_type, search_term=""):
    scope = f"Generated from {module_label} with {rows_count} row{'s' if rows_count != 1 else ''}."
    filter_note = f" Search filter: {search_term}." if search_term else ""
    type_note = f" Document type: {humanize_token(document_type)}."
    return f"{scope}{type_note}{filter_note}".strip()


def build_document_entries(rows):
    entries = []
    for index, row in enumerate(rows[:50], start=1):
        items = []
        for key, value in row.items():
            if key in EXCLUDED_ROW_KEYS:
                continue
            formatted = format_scalar(value)
            if not formatted:
                continue
            items.append((humanize_token(key), formatted))
        entries.append({"title": choose_row_title(row, index), "items": items})
    return entries


def build_markdown_document(*, title, module_label, report_preset, document_type, rows, search_term=""):
    summary = build_document_summary(module_label, len(rows), document_type, search_term)
    lines = [
        f"# {title}",
        "",
        f"- Module: {module_label}",
        f"- Report view: {humanize_token(report_preset or 'report')}",
        f"- Document type: {humanize_token(document_type)}",
        f"- Record count: {len(rows)}",
    ]
    if search_term:
        lines.append(f"- Search filter: {search_term}")

    lines.extend(["", "## Executive Summary", "", summary, "", "## Detailed Entries", ""])

    for entry in build_document_entries(rows):
        lines.append(f"### {entry['title']}")
        lines.append("")
        for label, value in entry["items"]:
            lines.append(f"- {label}: {value}")
        lines.append("")

    if len(rows) > 50:
        lines.extend(
            [
                "## Note",
                "",
                f"The snapshot contains {len(rows)} rows. This generated document shows the first 50 entries for readability.",
                "",
            ]
        )

    return "\n".join(lines).strip() + "\n"


def build_json_snapshot(*, title, module_label, report_preset, document_type, rows, search_term=""):
    return json.dumps(
        {
            "title": title,
            "module_label": module_label,
            "report_preset": report_preset,
            "document_type": document_type,
            "search_term": search_term,
            "rows": rows,
        },
        ensure_ascii=False,
        indent=2,
        default=str,
    )


def build_text_document(*, title, module_label, report_preset, document_type, rows, search_term=""):
    return build_markdown_document(
        title=title,
        module_label=module_label,
        report_preset=report_preset,
        document_type=document_type,
        rows=rows,
        search_term=search_term,
    ).replace("# ", "").replace("## ", "").replace("### ", "")


def build_document_payload(*, title, module_label, report_preset, document_type, output_format, rows, search_term=""):
    markdown = build_markdown_document(
        title=title,
        module_label=module_label,
        report_preset=report_preset,
        document_type=document_type,
        rows=rows,
        search_term=search_term,
    )
    entries = build_document_entries(rows)
    summary = build_document_summary(module_label, len(rows), document_type, search_term)

    if output_format == "json":
        content_text = build_json_snapshot(
            title=title,
            module_label=module_label,
            report_preset=report_preset,
            document_type=document_type,
            rows=rows,
            search_term=search_term,
        )
        return {"content_text": content_text, "file_bytes": content_text.encode("utf-8"), "mime_type": TEXT_CONTENT_TYPES["json"]}

    if output_format == "text":
        content_text = build_text_document(
            title=title,
            module_label=module_label,
            report_preset=report_preset,
            document_type=document_type,
            rows=rows,
            search_term=search_term,
        )
        return {"content_text": content_text, "file_bytes": content_text.encode("utf-8"), "mime_type": TEXT_CONTENT_TYPES["text"]}

    if output_format == "pdf":
        return {
            "content_text": markdown,
            "file_bytes": render_pdf_document(title=title, module_label=module_label, report_preset=report_preset, document_type=document_type, summary=summary, entries=entries, rows_count=len(rows), search_term=search_term),
            "mime_type": TEXT_CONTENT_TYPES["pdf"],
        }

    if output_format == "docx":
        return {
            "content_text": markdown,
            "file_bytes": render_docx_document(title=title, module_label=module_label, report_preset=report_preset, document_type=document_type, summary=summary, entries=entries, rows_count=len(rows), search_term=search_term),
            "mime_type": TEXT_CONTENT_TYPES["docx"],
        }

    return {"content_text": markdown, "file_bytes": markdown.encode("utf-8"), "mime_type": TEXT_CONTENT_TYPES["markdown"]}


def render_pdf_document(*, title, module_label, report_preset, document_type, summary, entries, rows_count, search_term=""):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("PDF export requires the reportlab package.") from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="MutedMeta", parent=styles["BodyText"], fontSize=9, leading=12, textColor=colors.HexColor("#555555")))
    styles.add(ParagraphStyle(name="EntryHeading", parent=styles["Heading2"], fontSize=13, leading=16, textColor=colors.HexColor("#111111"), spaceAfter=6))
    styles["Title"].fontSize = 22
    styles["Title"].leading = 26

    story = [
        Paragraph(title, styles["Title"]),
        Spacer(1, 6),
        Paragraph(f"Module: {module_label}", styles["MutedMeta"]),
        Paragraph(f"Report view: {humanize_token(report_preset or 'report')}", styles["MutedMeta"]),
        Paragraph(f"Document type: {humanize_token(document_type)}", styles["MutedMeta"]),
        Paragraph(f"Record count: {rows_count}", styles["MutedMeta"]),
    ]
    if search_term:
        story.append(Paragraph(f"Search filter: {search_term}", styles["MutedMeta"]))

    story.extend([Spacer(1, 14), Paragraph("Executive Summary", styles["Heading2"]), Paragraph(summary, styles["BodyText"]), Spacer(1, 14)])

    for entry in entries:
        story.append(Paragraph(entry["title"], styles["EntryHeading"]))
        table_data = [[label, value] for label, value in entry["items"]] or [["Details", "No additional fields available."]]
        table = Table(table_data, colWidths=[50 * mm, 110 * mm], repeatRows=0)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#111111")),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("LEADING", (0, 0), (-1, -1), 11),
                    ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D4CCC2")),
                    ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#D4CCC2")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.extend([table, Spacer(1, 10)])

    if rows_count > 50:
        story.extend([Paragraph("Note", styles["Heading3"]), Paragraph(f"The full snapshot contains {rows_count} rows. This export includes the first 50 entries for readability.", styles["BodyText"])])

    doc.build(story)
    return buffer.getvalue()


def render_docx_document(*, title, module_label, report_preset, document_type, summary, entries, rows_count, search_term=""):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("DOCX export requires the python-docx package.") from exc

    document = Document()
    document.core_properties.title = title

    title_paragraph = document.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta_lines = [
        f"Module: {module_label}",
        f"Report view: {humanize_token(report_preset or 'report')}",
        f"Document type: {humanize_token(document_type)}",
        f"Record count: {rows_count}",
    ]
    if search_term:
        meta_lines.append(f"Search filter: {search_term}")

    for line in meta_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    document.add_paragraph()
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(summary)

    for entry in entries:
        document.add_heading(entry["title"], level=2)
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        header_cells[0].text = "Field"
        header_cells[1].text = "Value"

        for label, value in entry["items"] or [("Details", "No additional fields available.")]:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        document.add_paragraph()

    if rows_count > 50:
        document.add_heading("Note", level=2)
        document.add_paragraph(f"The full snapshot contains {rows_count} rows. This export includes the first 50 entries for readability.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
